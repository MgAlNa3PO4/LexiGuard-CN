# -*- coding: utf-8 -*-
"""
网络言论法律分析 Agent

流程：
  1. 多角度拆解：把待分析文本拆成 3-5 个检索角度（DeepSeek）
  2. 分角度检索：每个角度 search() 取 top5，合并去重得到候选法条池
  3. 大模型判断：原文 + 候选法条池 → DeepSeek 判定是否违法/风险等级/适用条款/行为描述/处理建议
  4. 生成 Word 报告：固定格式 .docx，保存到 reports/ 目录

对外接口：
  from analyzer import analyze, generate_report
  result = analyze(text)
  path = generate_report(text, result)
"""

import os
import re
import json
import datetime

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from search import search, search_cases

# ============================================================================
# ★★★ DeepSeek API 配置 —— 在这里填入你的 API Key ★★★
# 两种方式任选其一：
#   (1) 直接把 key 写到下面引号里： DEEPSEEK_API_KEY = "sk-xxxxxxxx"
#   (2) 设置环境变量： set DEEPSEEK_API_KEY=sk-xxxxxxxx  （代码会自动读取）
# ============================================================================
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")  # <-- 在此填入你的 key，或设置环境变量 DEEPSEEK_API_KEY

DEEPSEEK_BASE_URL = "https://api.deepseek.com/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"

# 仅依赖本文件所在目录，文件夹改名也无需改代码
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports")

PER_ANGLE_TOPK = 8  # 每个检索角度取前几条（调大以提升诽谤等条款的召回完整性）


# ----------------------------------------------------------------------------
# DeepSeek 调用
# ----------------------------------------------------------------------------
def _deepseek_chat(messages, temperature=0.2, json_mode=True, timeout=60):
    """调用 DeepSeek chat 接口，返回 assistant 文本内容。"""
    if not DEEPSEEK_API_KEY:
        raise RuntimeError(
            "未配置 DeepSeek API Key。请在 analyzer.py 顶部的 DEEPSEEK_API_KEY 处填入，"
            "或设置环境变量 DEEPSEEK_API_KEY。"
        )
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": messages,
        "temperature": temperature,
        "stream": False,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}

    resp = requests.post(DEEPSEEK_BASE_URL, headers=headers, json=payload, timeout=timeout)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def _parse_json(text):
    """从模型输出中稳健解析 JSON（兼容 ```json 代码块包裹的情况）。"""
    text = text.strip()
    m = re.search(r"```(?:json)?\s*(\{.*\}|\[.*\])\s*```", text, re.S)
    if m:
        text = m.group(1)
    else:
        m = re.search(r"(\{.*\}|\[.*\])", text, re.S)
        if m:
            text = m.group(1)
    return json.loads(text)


# 平台中性化：确保模型输出里不出现“微博”这一平台专有词（报告通用化）
_NEUTRALIZE_RULES = [
    # 先匹配较长的复合词，避免“该微博内容”被替成“该该内容”
    ("该微博内容", "该内容"),
    ("本微博内容", "该内容"),
    ("微博发布者", "发布者"),
    ("微博内容", "该内容"),
    ("微博文本", "该内容"),
    ("该微博", "该内容"),
    ("本微博", "该内容"),
    ("微博平台", "网络平台"),
    ("微博", "网络言论"),
]


def _neutralize(text):
    if not text:
        return text
    for a, b in _NEUTRALIZE_RULES:
        text = text.replace(a, b)
    return text


# ----------------------------------------------------------------------------
# 第一步：多角度拆解
# ----------------------------------------------------------------------------
def decompose_query(text, n_min=3, n_max=5):
    """把待分析文本拆成 3-5 个用于法条检索的角度短语。"""
    sys_prompt = (
        "你是网络言论法律分析助手。请把给定的网络言论文本，从【可能涉及的违法/违规风险】角度，"
        f"拆解成 {n_min}-{n_max} 个简短的检索角度短语（每个 4-12 字），用于到法律条文库里检索相关法条。"
        "角度要覆盖文本可能触及的不同法律风险点，例如：散布谣言、侮辱他人名誉、煽动闹事、"
        "侵犯隐私、寻衅滋事、妨碍公务等。"
        "特别注意：要明确区分【发布者本人的行为】和【发布者所描述/评论的他人的行为】，"
        "只针对发布者自己发表的言论提取法律风险角度，"
        "不要把发布者评论的对象（第三方）的行为当成法律风险点。"
        '只输出 JSON，格式：{"angles": ["角度1", "角度2", ...]}'
    )
    content = _deepseek_chat(
        [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": f"网络言论文本：\n{text}"},
        ],
        temperature=0.3,
    )
    data = _parse_json(content)
    angles = [a.strip() for a in data.get("angles", []) if a and a.strip()]
    if not angles:
        angles = [text[:30]]  # 兜底：用原文片段作为单一角度
    return angles[:n_max]


# ----------------------------------------------------------------------------
# 第二步：分角度检索 + 合并去重
# ----------------------------------------------------------------------------
def retrieve_candidates(angles, per_angle_topk=PER_ANGLE_TOPK):
    """每个角度检索 top_k，按《法律名+条号》去重，保留命中的检索角度与最高分。"""
    pool = {}
    for angle in angles:
        for r in search(angle, top_k=per_angle_topk):
            key = (r["law"], r["article"])
            if key not in pool:
                r = dict(r)
                r["hit_angles"] = [angle]
                pool[key] = r
            else:
                if angle not in pool[key]["hit_angles"]:
                    pool[key]["hit_angles"].append(angle)
                pool[key]["rrf_score"] = max(pool[key]["rrf_score"], r["rrf_score"])
    # 按 RRF 分数排序
    return sorted(pool.values(), key=lambda x: x["rrf_score"], reverse=True)


# ----------------------------------------------------------------------------
# 第三步：大模型判断
# ----------------------------------------------------------------------------
def judge(text, candidates, case_candidates=None):
    """把原文 + 候选法条池 + 候选判例交给 DeepSeek，输出结构化判定。"""
    case_candidates = case_candidates or []
    # 候选法条编号列表（供模型按 index 引用，避免编造条文）
    lines = []
    for i, c in enumerate(candidates):
        lines.append(f"[{i}] 《{c['law']}》{c['article']}：{c['content']}")
    candidate_block = "\n".join(lines)

    # 候选判例编号列表（供模型按 index 引用，避免编造案例/判决）
    case_lines = []
    for i, c in enumerate(case_candidates):
        case_lines.append(
            f"[{i}] {c.get('case_name', '')}（相似度{c.get('similarity', '?')}%）"
            f"｜言论类型：{c.get('speech_type', '')}"
            f"｜涉案言论：{(c.get('speech_content') or '')[:60]}"
            f"｜法院认定：{c.get('court_finding', '')}"
            f"｜判决结果：{c.get('verdict', '')}"
            f"｜关键判断标准：{(c.get('key_standard') or '')[:80]}"
        )
    case_block = "\n".join(case_lines) if case_lines else "（无候选判例）"

    sys_prompt = (
        "你是专业的网络言论法律分析助手，只判断该内容发布者本人的言论行为是否违法，不分析内容中提及的第三方行为。\n"
        "下面给你一段网络言论文本，以及从法律条文库检索出的候选法条（带编号）。"
        "请你判断该内容是否涉嫌违法违规，并【只】从候选法条中挑选真正相关的条款（不要编造、不要引用未给出的条文）。"
        "判断要客观、留有余地：证据不足时用“存疑”，不要过度定性。\n"
        "对可能升级为治安或刑事责任的情形，即使当前存疑也列出对应条款，"
        "注明当前为何不够门槛、以及什么情况下会构成违法。\n"
        "\n请按以下【三层判断标准】定性，不要一刀切，要找准中间点：\n"
        "\n第一层 —— 明确不违法（is_illegal=否，risk_level=无）：\n"
        "- 对公众人物（含公职人员）行为、决策的主观评价；\n"
        "- 基于公开事实的舆论监督；\n"
        "- 对公共事件表达不满。\n"
        "  法律依据：《中华人民共和国民法典》第一千零二十五条——行为人为公共利益实施新闻报道、"
        "舆论监督等行为，影响他人名誉的，不承担民事责任。此类内容判“否/无”，不提取话语证据。\n"
        "\n第二层 —— 存疑需标注（is_illegal=存疑，risk_level=低）：\n"
        "- 对具名个人使用带有侮辱色彩的词汇（如“不配”“不配执掌”“愚弄”“愚弄民意”“护短纵容”等），"
        "即使针对公职人员、即使属于批评监督，只要措辞激烈带人身贬损色彩，也需标注；\n"
        "- 无明确证据的严重指控（如“形成密不透风的特权闭环”这类未经证实的定性指控）。\n"
        "  此类内容判“存疑/低”，并且【必须】把相关原话提取为话语证据。\n"
        "\n第三层 —— 明确违法（is_illegal=是，risk_level=中或高）：\n"
        "- 捏造具体、可核查的虚假事实；\n"
        "- 散布谣言并已造成传播；\n"
        "- 直接使用侮辱性词汇攻击普通个人（非公众人物）。\n"
        "  造成大规模传播、或捏造事实情节严重的判“高”，否则判“中”，并【必须】提取话语证据。\n"
        "\n【公开信/批评公职人员类文章的尺度】对公职人员的批评属于舆论监督，容忍度更高，"
        "不要轻易认定违法；但“不配执掌”“愚弄民意”这类对具名个人带侮辱色彩的措辞，"
        "应提取为证据并标注为低风险，而不是完全忽略。存疑时风险默认取“低”、不拔高到“中”，"
        "但也不要因整体偏向监督就把这些措辞全部漏掉。\n"
        "只输出 JSON，字段如下：\n"
        '{\n'
        '  "is_illegal": "是 | 否 | 存疑",\n'
        '  "risk_level": "高 | 中 | 低 | 无",\n'
        '  "articles": [{"index": 候选法条编号(整数), "reason": "为什么这条适用于本文"}],\n'
        '  "behavior": "具体说明该内容的哪些部分构成或涉嫌违法",\n'
        '  "suggestion": "针对性的处置建议",\n'
        '  "evidence_mapping": [{"quote": "原文中的关键违规话语，直接引用、不改写，10-50字", '
        '"index": 对应的候选法条编号(整数), '
        '"reason": "一句话说明这句话为何涉嫌违反该条款"}],\n'
        '  "escalation": [{"index": 候选法条编号(整数), '
        '"condition": "升级为违法/犯罪所需满足的条件，以及当前为何不够门槛", '
        '"suggestion": "针对该升级风险的监控/处置建议"}],\n'
        '  "reference_cases": [{"index": 候选判例编号(整数), '
        '"relevance": "为什么参考这个案例：本内容与该案在言论类型/传播规模/认定标准上的相似点，'
        '以及可参考该案的什么判断标准"}]\n'
        '}\n'
        "【话语证据提取规则（重要）】articles 与 evidence_mapping 用于承载第二层、第三层的内容：\n"
        "- 若整体为第一层（is_illegal=否，纯粹的主观评价、对公众人物决策的评价、正常舆论监督），"
        "articles 与 evidence_mapping 均为空数组，不把受保护的观点表达列为违规证据。\n"
        "- 只要文中存在第二层或第三层的内容，就【必须】把相关原话提取到 evidence_mapping，"
        "不能因为整体定性为“存疑”就不提取证据。\n"
        "- 每条证据的 reason 必须如实标注该句的风险级别，以“【低风险】”“【中风险】”“【高风险】”开头，"
        "再用一句话说明涉嫌违反的理由；属于第二层（存疑/低风险）的，理由用“带有侮辱色彩”"
        "“可能构成”“涉嫌”等审慎表述，不要武断定性。\n"
        "evidence_mapping 从原文中挑选最多 5 句最关键、确实属于第二/三层的原话（没有则留空数组，不凑数）；"
        "每句 quote 必须是原文的直接引用（逐字照抄，不得改写、不得拼接多句），长度控制在 10-50 字；"
        "index 必须来自上面候选法条编号，每句话只对应一条最相关的法律（一对一，严禁一句对多条）。\n"
        "【articles 违法说明的撰写要求】每一条涉及法条的违法说明(reason)，必须对应 evidence_mapping 中"
        "一句【不同的、具体的】原文 quote 作为证据支撑：不同法条要引用不同的原文句子，"
        "严禁多条法条重复引用同一句话、或套用同一句模板化表述。"
        "若 evidence_mapping 中的 quote 数量少于 articles 的数量，则优先为风险最高的法条配上具体 quote 证据；"
        "其余没有专属 quote 的法条，其违法说明改为阐述该条法律【适用于本内容的整体性原因】"
        "（例如内容的整体性质、表述方式、传播范围或可能造成的后果），而不是重复引用同一句原文。\n"
        "escalation 用于列出当前虽不构成、但可能升级为治安或刑事责任的条款。"
        "其 index 必须来自上面给定的候选法条编号（条款与条号一律以候选法条为准，"
        "严禁自行编造或凭记忆填写法律名称、条号）；若候选法条中没有合适的升级条款，escalation 用空数组。"
        "condition 与 suggestion 只写结论性内容，不要写自我推敲、反问或对条号是否准确的说明。\n"
        "【参考判例（reference_cases）撰写要求】下方会给出从最高法网络暴力典型案例库检索出的候选判例（带编号）。"
        "请从中挑选与本内容【确有参照价值】的判例（最多 3 个，按相似度从高到低），其 index 必须来自候选判例编号，"
        "严禁编造案例名称或判决结果。relevance 要具体说明本内容与该案在【言论类型、传播规模、关键认定标准】上的相似点，"
        "以及可参考该案的什么标准（例如传播规模的认定、是否‘情节严重’、是否‘严重危害社会秩序’等）。"
        "若本内容属于第一层（合法的舆论监督/对公众人物的主观评价），可只保留用于划清边界的判例并在 relevance 中说明二者区别，"
        "或在确无可参照判例时返回空数组。不要为了凑数而牵强附会。"
    )
    user_prompt = (
        f"【待分析内容】\n{text}\n\n"
        f"【候选法条】\n{candidate_block}\n\n"
        f"【候选判例】\n{case_block}\n\n"
        "请输出 JSON 判定结果。"
    )
    content = _deepseek_chat(
        [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
    )
    data = _parse_json(content)

    # 把模型引用的 index 映射回完整法条
    matched = []
    for item in data.get("articles", []):
        try:
            idx = int(item.get("index"))
        except (TypeError, ValueError):
            continue
        if 0 <= idx < len(candidates):
            c = candidates[idx]
            matched.append(
                {
                    "law": c["law"],
                    "article": c["article"],
                    "content": c["content"],
                    "reason": _neutralize(item.get("reason", "").strip()),
                }
            )

    # 风险升级提示：条款按 index 从候选法条池引用，条号/条文一律取自知识库
    escalation = []
    for item in data.get("escalation", []) or []:
        try:
            idx = int(item.get("index"))
        except (TypeError, ValueError):
            continue
        if 0 <= idx < len(candidates):
            c = candidates[idx]
            escalation.append(
                {
                    "law": c["law"],
                    "article": c["article"],
                    "content": c["content"],
                    "condition": _neutralize((item.get("condition") or "").strip()),
                    "suggestion": _neutralize((item.get("suggestion") or "").strip()),
                }
            )

    # 话语证据与法条对应：quote/reason 由模型给出，法律名/条号/条文按 index 取自知识库
    evidence_mapping = []
    for item in data.get("evidence_mapping", []) or []:
        quote = (item.get("quote") or "").strip().strip("“”\"")
        try:
            idx = int(item.get("index"))
        except (TypeError, ValueError):
            continue
        if quote and 0 <= idx < len(candidates):
            c = candidates[idx]
            evidence_mapping.append(
                {
                    "quote": quote,
                    "law": c["law"],
                    "article": c["article"],
                    "content": c["content"],
                    "reason": _neutralize((item.get("reason") or "").strip()),
                }
            )

    # 参考判例：relevance 由模型给出，案例名/相似度/言论/判决一律按 index 取自判例库
    reference_cases = []
    for item in data.get("reference_cases", []) or []:
        try:
            idx = int(item.get("index"))
        except (TypeError, ValueError):
            continue
        if 0 <= idx < len(case_candidates):
            c = case_candidates[idx]
            reference_cases.append(
                {
                    "case_name": c.get("case_name", ""),
                    "similarity": c.get("similarity"),
                    "speech": (c.get("speech_content") or "").strip(),
                    "verdict": c.get("verdict", ""),
                    "relevance": _neutralize((item.get("relevance") or "").strip()),
                }
            )

    is_illegal = data.get("is_illegal", "存疑")
    risk_level = data.get("risk_level", "无")

    # 一致性兜底：判定为“否”（不违法，如正常舆论监督/对公众人物的主观评价）时，
    # 不应再出现“涉嫌违反某条”的法条与话语证据——避免结论与证据自相矛盾。
    # 相关风险（若有）一律保留在 escalation（风险预警）中呈现。
    if is_illegal == "否":
        matched = []
        evidence_mapping = []
        if risk_level not in ("无", "低"):
            risk_level = "低"

    return {
        "is_illegal": is_illegal,
        "risk_level": risk_level,
        "articles": matched,
        "behavior": _neutralize(data.get("behavior", "").strip()),
        "suggestion": _neutralize(data.get("suggestion", "").strip()),
        "evidence_mapping": evidence_mapping,
        "escalation": escalation,
        "reference_cases": reference_cases,
    }


# ----------------------------------------------------------------------------
# 主编排
# ----------------------------------------------------------------------------
def analyze(text):
    """完整分析流程，返回结果 dict。"""
    text = (text or "").strip()
    if not text:
        raise ValueError("待分析文本为空。")

    angles = decompose_query(text)
    candidates = retrieve_candidates(angles)
    # 判例检索：用原文到判例库召回最相似的典型案例，交给 judge 选择并说明参考意义
    try:
        case_candidates = search_cases(text, top_k=3)
    except Exception:
        case_candidates = []
    verdict = judge(text, candidates, case_candidates)

    return {
        "text": text,
        "angles": angles,
        "candidate_count": len(candidates),
        **verdict,
    }


# ----------------------------------------------------------------------------
# 第四步：生成 Word 报告
# ----------------------------------------------------------------------------
SEP = "━" * 26
FONT_NAME = "宋体"


def _set_song(run, size, bold=False):
    """把 run 设为宋体（同时设置中文 eastAsia 字体），指定字号与加粗。"""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    # run.font.name 只设置西文字体，中文需单独设置 eastAsia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _para(doc, text, size=12, bold=False, align=None):
    """新增一个宋体段落。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_song(run, size, bold)
    return p


def _brief(content, n=45):
    """把完整条文压成简述：去掉开头条号，截断到 n 字。"""
    s = re.sub(r"^第[^条]*条(?:之[一二三四五六七八九十]+)?\s*", "", content).strip()
    return s[:n] + ("……" if len(s) > n else "")


def _add_sep(doc):
    # 分割线保留不变（仍用宋体渲染以保持整体一致）
    p = _para(doc, SEP, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    return p


def _add_heading_line(doc, text):
    # 模块标题：宋体 12pt 加粗
    return _para(doc, text, size=12, bold=True)


def generate_report(text, result, out_dir=REPORTS_DIR, excerpt=False, source=None):
    """把分析结果写成固定格式的 .docx，返回保存路径。

    excerpt: 为 True 时在原文内容上方注明“以下为文章节选分析”。
    source:  可选 dict，公众号文章来源信息 {title, account, url}。
    """
    os.makedirs(out_dir, exist_ok=True)
    now = datetime.datetime.now()
    fname = "言论检测报告_" + now.strftime("%Y%m%d%H%M") + ".docx"
    path = os.path.join(out_dir, fname)

    doc = Document()
    # 设置全局默认字体为宋体 12pt（兜底，保证未显式设置的文本也是宋体）
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # 标题
    _add_sep(doc)
    # 报告主标题：宋体 16pt 加粗 居中
    _para(doc, "网络言论法律分析报告", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    # 分析时间：宋体 11pt 居中
    _para(doc, "分析时间：" + now.strftime("%Y年%m月%d日 %H:%M"),
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_sep(doc)

    # 原文内容
    _add_heading_line(doc, "【原文内容】")
    if source:
        meta_bits = []
        if source.get("title"):
            meta_bits.append("文章标题：" + source["title"])
        if source.get("account"):
            meta_bits.append("公众号：" + source["account"])
        if source.get("url"):
            meta_bits.append("来源链接：" + source["url"])
        for b in meta_bits:
            _para(doc, b, size=12)
        if meta_bits:
            _para(doc, "")
    if excerpt:
        _para(doc, "（以下为文章节选分析，仅截取前 3000 字）", size=12, bold=True)
    _para(doc, result["text"], size=12)
    _para(doc, "")

    # 1. 分析结论
    _add_heading_line(doc, "【分析结论】")
    _para(doc, "是否违法：" + result.get("is_illegal", "存疑"), size=12)
    _para(doc, "风险等级：" + result.get("risk_level", "无"), size=12)
    _para(doc, "")

    # 2. 话语证据与法条对应（移到最前）
    _add_sep(doc)
    _add_heading_line(doc, "【话语证据与法条对应】")
    _para(doc, "")
    evidence = result.get("evidence_mapping", [])
    if not evidence:
        _para(doc, "（未提取到明确的话语证据）", size=12)
    else:
        for i, e in enumerate(evidence, 1):
            _para(doc, f"证据{i}：", size=12, bold=True)
            _para(doc, "“" + e["quote"] + "”", size=12)
            _para(doc, "↓ 涉嫌违反", size=12)
            _para(doc, "《{}》{}：{}".format(e["law"], e["article"], _brief(e["content"])), size=12)
            _para(doc, "说明：" + (e.get("reason") or "（无）"), size=12)
            _para(doc, "")
    _add_sep(doc)

    # 3. 涉及法条
    _add_heading_line(doc, "【涉及法条】")
    articles = result.get("articles", [])
    if not articles:
        _para(doc, "（未匹配到明确适用的法律条款）", size=12)
    else:
        for i, a in enumerate(articles, 1):
            _para(doc, f"{i}. 《{a['law']}》{a['article']}", size=12, bold=True)
            _para(doc, "   条文内容：" + a["content"], size=12)
            _para(doc, "   违法说明：" + (a.get("reason") or "（无）"), size=12)
            _para(doc, "")

    # 4. 参考判例（放在涉及法条之后）
    _para(doc, "")
    _add_heading_line(doc, "【参考判例】")
    reference_cases = result.get("reference_cases", [])
    if not reference_cases:
        _para(doc, "（未匹配到可参考的典型判例）", size=12)
    else:
        for i, rc in enumerate(reference_cases, 1):
            sim = rc.get("similarity")
            sim_txt = f"（相似度{sim}%）" if sim is not None else ""
            _para(doc, f"案例{i}：{rc.get('case_name', '')}{sim_txt}", size=12, bold=True)
            speech = (rc.get("speech") or "").strip()
            if len(speech) > 60:
                speech = speech[:60] + "……"
            _para(doc, "涉案言论：“" + speech + "”", size=12)
            _para(doc, "判决结果：" + (rc.get("verdict") or "（无）"), size=12)
            _para(doc, "参考意义：" + (rc.get("relevance") or "（无）"), size=12)
            _para(doc, "")
    _add_sep(doc)

    # 5. 处理建议与风险预警（合并）
    _add_heading_line(doc, "【处理建议与风险预警】")
    _para(doc, "当前建议：" + (result.get("suggestion") or "（无）"), size=12)
    _para(doc, "─" * 24, size=12)  # 细分隔线
    escalation = result.get("escalation", [])
    if not escalation:
        _para(doc, "风险预警：当前未发现明显的升级风险。", size=12)
    else:
        _para(doc, "风险预警：当前定性为{}/{}风险，若出现以下情况可能升级：".format(
            result.get("is_illegal", "存疑"), result.get("risk_level", "低")), size=12)
        for e in escalation:
            _para(doc, "- 涉及条款：《{}》{}".format(e["law"], e["article"]), size=12)
            _para(doc, "  条文内容：" + e["content"], size=12)
            _para(doc, "  升级条件：" + (e.get("condition") or "（无）"), size=12)
            _para(doc, "  建议：" + (e.get("suggestion") or "（无）"), size=12)
            _para(doc, "")

    _add_sep(doc)
    doc.save(path)
    return path


if __name__ == "__main__":
    demo = "这家公司的产品全是假货，老板就是个骗子，大家千万别买"
    res = analyze(demo)
    print(json.dumps(res, ensure_ascii=False, indent=2))
    print("报告已生成：", generate_report(demo, res))
